import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

src = open('SF4030T-5S_RU.md', encoding='utf-8').read().split('\n')

doc = Document()
for s in ('Normal',):
    st = doc.styles[s]
    st.font.name = 'Times New Roman'
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
doc.styles['Normal'].paragraph_format.space_after = Pt(6)
doc.styles['Normal'].paragraph_format.line_spacing = 1.15

sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.0)
sec.top_margin = sec.bottom_margin = Cm(1.8)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:shd'); el.set(qn('w:fill'), color); tcPr.append(el)

INL = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]*?\))')
def runs(p, text):
    text = text.replace('<br>', '\n')
    for part in INL.split(text):
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Consolas'
        elif part.startswith('['):
            m = re.match(r'\[([^\]]+)\]\([^)]*\)', part)
            p.add_run(m.group(1) if m else part)
        else:
            for i, chunk in enumerate(part.split('\n')):
                if i: p.add_run().add_break()
                p.add_run(chunk)

i = 0
n = len(src)
first_h1 = True
while i < n:
    line = src[i].rstrip()
    # рисунок
    mimg = re.match(r'^!\[(.*?)\]\((fig/[^)]+)\)\s*$', line)
    if mimg:
        from PIL import Image as PILImage
        cap, path = mimg.group(1), mimg.group(2)
        px_w, px_h = PILImage.open(path).size
        dpi = 300 if path.rsplit('/', 1)[1] in ('sign_shock.png', 'label_lub.png') else 260
        w_cm = px_w / dpi * 2.54
        w_cm = min(w_cm, 16.0)
        ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.paragraph_format.space_before = Pt(6); ip.paragraph_format.space_after = Pt(2)
        ip.add_run().add_picture(path, width=Cm(w_cm))
        if cap:
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
            runs(cp, cap)
            for r in cp.runs:
                r.italic = True; r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        i += 1
        continue

    if line.startswith('|') and i + 1 < n and re.match(r'^\|[\s:|-]+\|$', src[i+1].strip()):
        rows = []
        while i < n and src[i].strip().startswith('|'):
            cells = [c.strip() for c in src[i].strip().strip('|').split('|')]
            if not re.match(r'^[\s:|-]+$', src[i].strip().strip('|')):
                rows.append(cells)
            i += 1
        cols = max(len(r) for r in rows)
        t = doc.add_table(rows=0, cols=cols)
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            cs = t.add_row().cells
            for ci in range(cols):
                txt = row[ci] if ci < len(row) else ''
                para = cs[ci].paragraphs[0]
                para.paragraph_format.space_after = Pt(2)
                runs(para, txt)
                for r in para.runs:
                    r.font.size = Pt(9)
                    if ri == 0: r.bold = True
                if ri == 0: shade(cs[ci], 'E6E6E6')
        doc.add_paragraph()
        continue
    if line.startswith('---') and set(line.strip()) == {'-'}:
        i += 1; continue
    if line.startswith('#'):
        lvl = len(line) - len(line.lstrip('#'))
        txt = line.lstrip('#').strip()
        if lvl == 1 and not first_h1:
            doc.add_page_break()
        if lvl == 1: first_h1 = False
        h = doc.add_heading(level=min(lvl, 4))
        runs(h, txt)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
            r.font.name = 'Times New Roman'
        i += 1; continue
    if line.startswith('>'):
        buf = []
        while i < n and src[i].startswith('>'):
            buf.append(src[i].lstrip('>').strip()); i += 1
        text = ' '.join(x for x in buf if x)
        text = re.sub(r'^#+\s*', '', text)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        runs(p, text)
        for r in p.runs:
            r.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
        continue
    m = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)$', line)
    if m:
        indent = len(m.group(1)) // 3
        style = 'List Number' if m.group(2)[0].isdigit() else 'List Bullet'
        p = doc.add_paragraph(style=style)
        p.paragraph_format.left_indent = Cm(0.75 + 0.6 * indent)
        p.paragraph_format.space_after = Pt(2)
        runs(p, m.group(3))
        i += 1; continue
    if not line.strip():
        i += 1; continue
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    runs(p, line.strip())
    i += 1

doc.save('SF4030T-5S_Руководство_RU.docx')
print('ok')
